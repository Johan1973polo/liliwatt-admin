"""
Génère les contrats pré-remplis (Contrat Partenariat + Avenant 1)
depuis les modèles Drive, remplit avec les infos candidat,
et uploade dans le dossier Drive du candidat.
"""
import io
import json
import os
import base64
import re
from docx import Document
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
from google.oauth2.service_account import Credentials


MODELES_FOLDER_ID = '1VZ1Fdrm-pfHF10OvYfGXCVEbhijAjtPy'


def get_drive_service():
    creds_b64 = os.environ.get('GOOGLE_DRIVE_CREDS_BASE64', '')
    creds_json = os.environ.get('GOOGLE_CREDS_JSON', '')
    if creds_b64:
        creds_dict = json.loads(base64.b64decode(creds_b64).decode())
    elif creds_json:
        creds_dict = json.loads(creds_json)
    else:
        with open(os.path.join(os.path.dirname(__file__), 'liliwatt-eddcc0bc9e18.json')) as f:
            creds_dict = json.load(f)
    creds = Credentials.from_service_account_info(
        creds_dict, scopes=['https://www.googleapis.com/auth/drive']
    )
    return build('drive', 'v3', credentials=creds)


def download_docx(drive, file_id):
    """Télécharge un fichier Drive en mémoire."""
    request = drive.files().get_media(fileId=file_id)
    buf = io.BytesIO()
    downloader = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    buf.seek(0)
    return buf


def fill_document(doc_bytes, prenom, nom, siren, qualite):
    """Remplace les champs dans un Document Word."""
    doc = Document(doc_bytes)

    replacements = {
        'M. / Mme ___________________________': f'M./Mme {prenom} {nom}',
        'M./Mme ___________________________': f'M./Mme {prenom} {nom}',
        'M. / Mme _________________________': f'M./Mme {prenom} {nom}',
        'SIRET : ___________________________': f'SIRET : {siren}' if siren else 'SIRET : ___________________________',
        'SIRET : _________________________': f'SIRET : {siren}' if siren else 'SIRET : _________________________',
        'Qualité : _________________________': f'Qualité : {qualite}' if qualite else 'Qualité : _________________________',
        'Qualité : ___________________________': f'Qualité : {qualite}' if qualite else 'Qualité : ___________________________',
        'Nom : ___________________________': f'Nom : {prenom} {nom}',
        'Nom : _________________________': f'Nom : {prenom} {nom}',
    }

    # Remplacer dans les paragraphes
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)

    # Remplacer dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            for run in para.runs:
                                if old in run.text:
                                    run.text = run.text.replace(old, new)

    # Sauver en mémoire
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def generate_contrats(prenom, nom, siren, qualite, candidat_drive_folder_id):
    """
    Télécharge les modèles, les remplit et les uploade dans le dossier du candidat.
    Retourne la liste des fichiers créés.
    """
    drive = get_drive_service()

    # Lister les modèles dans le dossier
    result = drive.files().list(
        q=f"'{MODELES_FOLDER_ID}' in parents and trashed=false and mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'",
        fields='files(id, name)',
        supportsAllDrives=True,
        includeItemsFromAllDrives=True
    ).execute()
    modeles = result.get('files', [])

    if not modeles:
        # Chercher aussi les Google Docs et les exporter en docx
        result2 = drive.files().list(
            q=f"'{MODELES_FOLDER_ID}' in parents and trashed=false and mimeType='application/vnd.google-apps.document'",
            fields='files(id, name)',
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        modeles = result2.get('files', [])

    created_files = []
    nom_upper = nom.upper()

    for modele in modeles:
        print(f"  📄 Traitement modèle: {modele['name']}")

        # Télécharger
        if 'google-apps.document' in modele.get('mimeType', ''):
            # Export Google Doc en docx
            request = drive.files().export_media(fileId=modele['id'],
                mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            buf = io.BytesIO()
            downloader = MediaIoBaseDownload(buf, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            buf.seek(0)
        else:
            buf = download_docx(drive, modele['id'])

        # Remplir
        filled = fill_document(buf, prenom, nom_upper, siren, qualite)

        # Nom du fichier de sortie
        base_name = modele['name'].replace('.docx', '').strip()
        output_name = f"{base_name}_{nom_upper}_{prenom}.docx"

        # Uploader dans le dossier du candidat
        media = MediaIoBaseUpload(filled,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        uploaded = drive.files().create(
            body={
                'name': output_name,
                'parents': [candidat_drive_folder_id],
                'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            },
            media_body=media,
            fields='id, name, webViewLink',
            supportsAllDrives=True
        ).execute()

        created_files.append({
            'name': uploaded['name'],
            'id': uploaded['id'],
            'link': uploaded.get('webViewLink', '')
        })
        print(f"  ✅ Uploadé: {uploaded['name']}")

    return created_files


if __name__ == '__main__':
    # Test local
    import sys
    if len(sys.argv) >= 5:
        prenom, nom, siren, qualite, folder_id = sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4], sys.argv[5] if len(sys.argv) > 5 else ''
        files = generate_contrats(prenom, nom, siren, qualite, folder_id)
        for f in files:
            print(f"  {f['name']} → {f['link']}")
    else:
        print("Usage: python generate_contrats.py <prenom> <nom> <siren> <qualite> <folder_id>")
